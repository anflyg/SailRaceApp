from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SCREENSHOT_DIR = ROOT / "docs" / "manual" / "screenshots"
OUTPUT_DOCX = ROOT / "docs" / "Aster_Race_Anvandarmanual.docx"
OUTPUT_MD = ROOT / "docs" / "Aster_Race_Anvandarmanual.md"

NAVY = RGBColor(7, 31, 61)
CYAN = RGBColor(53, 194, 255)
ORANGE = RGBColor(255, 133, 0)
MAGENTA = RGBColor(223, 55, 189)


def add_title(document: Document, text: str, size: int, color: RGBColor = NAVY) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = color


def add_section_heading(document: Document, text: str, color: RGBColor = NAVY) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = color


def add_subheading(document: Document, text: str, color: RGBColor = NAVY) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = color


def add_bullets(document: Document, lines: list[str]) -> None:
    for line in lines:
        paragraph = document.add_paragraph(line, style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(3)


def add_image(document: Document, file_name: str, caption: str) -> None:
    image_path = SCREENSHOT_DIR / file_name
    if not image_path.exists():
        document.add_paragraph(f"[Saknar bild: {file_name}]")
        return
    document.add_picture(str(image_path), width=Inches(3.2))
    paragraph = document.paragraphs[-1]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_p = document.add_paragraph(caption)
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_p.runs[0].italic = True
    caption_p.runs[0].font.size = Pt(9)


def build_docx() -> None:
    document = Document()
    style = document.styles["Normal"]
    style.font.name = "Aptos"
    style.font.size = Pt(11)

    add_title(document, "Aster Race", 34, NAVY)
    add_title(document, "Användarmanual", 24, CYAN)
    document.add_paragraph()
    subtitle = document.add_paragraph("Start, bana, segling och analys för kappsegling")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(13)
    subtitle.runs[0].font.color.rgb = ORANGE
    document.add_page_break()

    add_section_heading(document, "1. Snabb översikt", CYAN)
    add_bullets(document, [
        "Setup: kontroll av GPS, sensorer, COG, R/S och Layline-inställningar.",
        "Bana: sätt A/B (startlinje), K1 (kryssmärke), L1 (länsmärke) och vindriktning.",
        "Start: nedräkning 5-1 min, TTL/BURN och GPS-precision inför start.",
        "Segling: fart, riktning, VMG Vind/VMG Bana och Layline-varning.",
        "Analys: racebibliotek, översikt, startanalys, grafer och data/export.",
    ])

    add_section_heading(document, "2. Setup", CYAN)
    add_image(document, "01_setup.png", "Setup-vyn")
    add_bullets(document, [
        "GPS visar status och precision.",
        "Fart och COG visar filtrerad GPS-data.",
        "Motion och Heading visar sensortillgänglighet.",
        "R/S visar rullning (R) och stampning (S).",
        "Kalibrera nolläge sparar aktuell R/S som referens.",
        "Layline På/Av aktiverar varningen i Segling.",
        "Alpha kan ställas 70-110°, default 90°, +/- ändrar 1° per tryck.",
    ])

    add_section_heading(document, "3. Bana", CYAN)
    add_image(document, "02_bana.png", "Bana-vyn")
    add_bullets(document, [
        "A och B definierar startlinjen.",
        "K1 är kryssmärket och mål för VMG Bana/Layline.",
        "L1 är länsmärke och referens för banaxel.",
        "Vindpilen används för att mäta/sätta vindriktning.",
        "Markörfärger visar GPS-kvalitet (bra/dålig).",
        "Rensa bana tar bort samtliga satta banobjekt.",
    ])

    add_section_heading(document, "4. Start", CYAN)
    add_image(document, "03_start_idle.png", "Start-vy, timer ej startad")
    add_image(document, "04_start_running.png", "Start-vy, timer igång")
    add_bullets(document, [
        "Välj nedräkning: 5, 4, 3, 2 eller 1 minut.",
        "Tryck på stora klockan för start/paus.",
        "Långtryck på klockan återställer timern.",
        "TTL = tid till startlinjen baserat på kurs/fart.",
        "BURN = tid att bränna (+) eller tid du saknar (-).",
        "GPS accuracy visas för snabb tillförlitlighetsbedömning.",
    ])

    add_section_heading(document, "5. Segling", CYAN)
    add_subheading(document, "5a. VMG Vind", CYAN)
    add_image(document, "06_segling_vmg_vind.png", "Segling med VMG Vind (cyan/blå)")
    add_subheading(document, "5b. VMG Bana", ORANGE)
    add_image(document, "05_segling_vmg_bana.png", "Segling med VMG Bana (orange)")
    add_subheading(document, "5c. Layline countdown", MAGENTA)
    add_image(document, "07_segling_layline.png", "Segling med Layline (lila/magenta)")
    add_bullets(document, [
        "Fart visar aktuell filtrerad SOG i knop.",
        "Riktning visar filtrerad COG.",
        "R/S längst ned visar rullning och stampning.",
        "VMG Vind = fartkomponent mot vindreferensen.",
        "VMG Bana = fartkomponent mot K1/banmålet.",
        "Tryck på VMG-rutan för att växla mellan VMG Vind och VMG Bana när båda finns.",
        "Layline visas ovanpå VMG-rutan och räknar 10 -> 0 -> -5.",
        "0 betyder appens bästa uppskattning av \"slå nu\".",
        "Ljudsignal spelas vid 10 och 0, tick spelas varje sekund 9-1.",
        "Layline använder K1, COG, SOG och Alpha, och är aktiv bara när båten seglar mot K1 enligt banreferens.",
    ])

    add_section_heading(document, "6. Analys", CYAN)
    add_image(document, "08_analys_bibliotek.png", "Analys: Racebibliotek")
    add_image(document, "09_analys_oversikt.png", "Analys: Översikt")
    add_image(document, "10_analys_startanalys.png", "Analys: Start")
    add_bullets(document, [
        "Racebibliotek visar sparade race, favoriter och grundstatistik.",
        "Översikt visar karta, replay-kontroller och sammanfattning.",
        "Start visar startanalys med resultat och osäkerhet.",
        "Grafer och Data är dedikerade sektioner för fördjupning.",
        "Race kan favoriteras, döpas om, exporteras och raderas.",
        "Layline-event sparas endast när aktiv race logging/sparad segling finns.",
    ])

    add_section_heading(document, "7. Kort arbetsflöde", CYAN)
    add_bullets(document, [
        "1) Setup: kontrollera sensorer, GPS och Alpha.",
        "2) Bana: sätt A, B, K1, L1 och vind.",
        "3) Start: välj timer och starta nedräkning.",
        "4) Segling: följ fart/riktning/VMG och Layline.",
        "5) Analys: granska och exportera sparad segling.",
    ])

    document.save(OUTPUT_DOCX)


def build_markdown() -> None:
    markdown = f"""# Aster Race - Användarmanual

_Start, bana, segling och analys för kappsegling_

## 1. Snabb översikt
- Setup: GPS, sensorer, COG, R/S och Layline-inställningar.
- Bana: A/B, K1, L1 och vindriktning.
- Start: timer, TTL/BURN och GPS accuracy.
- Segling: fart, riktning, VMG Vind/VMG Bana och Layline.
- Analys: bibliotek, översikt, startanalys, grafer, data.

## 2. Setup
![Setup](manual/screenshots/01_setup.png)
- GPS, Fart, COG, Motion, Heading, R/S och Kalibrera nolläge.
- Layline På/Av och Alpha 70-110° (default 90°, 1° steg med +/-).

## 3. Bana
![Bana](manual/screenshots/02_bana.png)
- A/B startlinje, K1 kryssmärke, L1 länsmärke/referens.
- Vindpil för mätning/sättning av vindriktning.
- Färgkodning visar GPS-kvalitet.

## 4. Start
![Start idle](manual/screenshots/03_start_idle.png)
![Start running](manual/screenshots/04_start_running.png)
- 5/4/3/2/1 minuter.
- Tryck stora klockan för start/paus, långtryck för reset.
- TTL = tid till linje, BURN = över-/undertid, samt GPS accuracy.

## 5. Segling
### 5a. VMG Vind
![VMG Vind](manual/screenshots/06_segling_vmg_vind.png)
### 5b. VMG Bana
![VMG Bana](manual/screenshots/05_segling_vmg_bana.png)
### 5c. Layline
![Layline](manual/screenshots/07_segling_layline.png)
- Fart och Riktning (COG) från filtrerad GPS.
- R/S nederst = rullning och stampning.
- VMG-rutan växlar mellan VMG Vind (cyan/blå) och VMG Bana (orange).
- Layline (magenta) ersätter VMG-rutan temporärt, räknar 10 till -5.
- 0 = bästa uppskattning av "slå nu"; ljud vid 10 och 0, tick 9-1.

## 6. Analys
![Bibliotek](manual/screenshots/08_analys_bibliotek.png)
![Översikt](manual/screenshots/09_analys_oversikt.png)
![Startanalys](manual/screenshots/10_analys_startanalys.png)
- Racebibliotek, Översikt, Start, Grafer och Data.
- Favorit, döp om, exportera, radera när tillgängligt.
- Layline-event sparas bara när race logging är aktiv.

## 7. Kort arbetsflöde
1. Setup: kontrollera sensorer och Alpha.
2. Bana: sätt A, B, K1, L1 och vind.
3. Start: välj timer och starta.
4. Segling: följ fart/riktning/VMG och Layline.
5. Analys: granska sparad segling.
"""
    OUTPUT_MD.write_text(markdown, encoding="utf-8")


def main() -> None:
    build_docx()
    build_markdown()
    print(f"Skapad: {OUTPUT_DOCX}")
    print(f"Skapad: {OUTPUT_MD}")


if __name__ == "__main__":
    main()
